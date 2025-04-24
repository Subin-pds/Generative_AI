# importing libraries
from docx import Document
import speech_recognition as sr
from pydub import AudioSegment
from pydub.silence import split_on_silence
import os
import re
import unicodedata
from moviepy import AudioFileClip
import pdfplumber  # PyMuPDF

# Function to extract text from PDF
def extract_pdf_to_markdown(pdf_path):
    with pdfplumber.open(pdf_path) as pdf:
        md_text = ""

        for page in pdf.pages:
            text = page.extract_text()
            if text:
                md_text += text + "\n\n"

            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    row_text = "| " + " | ".join(cell or "" for cell in row) + " |"
                    md_text += row_text + "\n"
                md_text += "\n"

        return md_text

# Extract text and tables from DOCX file
def extract_docx_to_markdown(docx_path):
    doc = Document(docx_path)
    md_lines = []

    for para in doc.paragraphs:
        style = para.style.name.lower()

        # Detect heading
        if 'heading' in style:
            level = ''.join(filter(str.isdigit, style)) or "1"
            md_lines.append(f"{'#' * int(level)} {para.text}")
        else:
            line = ""
            for run in para.runs:
                text = run.text
                if run.bold:
                    text = f"**{text}**"
                if run.italic:
                    text = f"*{text}*"
                line += text
            md_lines.append(line)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = "| " + " | ".join([cell.text.strip() for cell in row.cells]) + " |"
                md_lines.append(row_text)
            md_lines.append("")  # Add a newline between tables

    return "\n\n".join(md_lines)

# Extract text from TXT file
def extract_txt_to_markdown(txt_path):
    with open(txt_path, "r", encoding="utf-8") as f:
        return f.read()

'''
Speech to text
'''
# create a speech recognition object
r = sr.Recognizer()

# Function to recognize speech in the audio file so that we don't repeat ourselves in in other functions
def transcribe_audio(path):
    # use the audio file as the audio source
    with sr.AudioFile(path) as source:
        audio_listened = r.record(source)
        # try converting it to text
        text = r.recognize_google(audio_listened)
    return text

# a function that splits the audio file into chunks on silence and applies speech recognition
def extract_text_from_audio(path):
    """Splitting the large audio file into chunks and apply speech recognition on each of these chunks"""

    # If the file is mp3 converting it into wav format
    if path.lower().endswith(".mp3"):
        wav_path = path.replace(".mp3", ".wav")
        print("Converting MP3 to WAV using moviepy...")
        audio_clip = AudioFileClip(path) # convert mp3 to wav using moviepy
        audio_clip.write_audiofile(wav_path) # Save as WAV
        path = wav_path #Undate the path
        path = r"{}".format(path)
    # open the audio file using pydub
    sound = AudioSegment.from_file(path)
    # split audio sound where silence is 500 miliseconds or more and get chunks
    chunks = split_on_silence(sound,
        # experiment with this value for your target audio file
        min_silence_len = 300,
        # adjust this per requirement
        silence_thresh = sound.dBFS-14,
        # keep the silence for 1 second, adjustable as well
        keep_silence=500,
    )
    folder_name = "audio-chunks"

     # Define the subfolder path
    output_dir = os.path.join("Data", "Output")
    os.makedirs(output_dir, exist_ok=True)  # Create the folder if it doesn't exist

      # Define the markdown file path
    folder_name = os.path.join(output_dir, folder_name)

    # create a directory to store the audio chunks
    if not os.path.isdir(folder_name):
        os.mkdir(folder_name)
    whole_text = ""
    # process each chunk
    for i, audio_chunk in enumerate(chunks, start=1):
        # export audio chunk and save it in
        # the `folder_name` directory.
        chunk_filename = os.path.join(folder_name, f"chunk{i}.wav")
        audio_chunk.export(chunk_filename, format="wav")
        # recognize the chunk
        try:
            text = transcribe_audio(chunk_filename)
        except sr.UnknownValueError as e:
            print("Error:", str(e))
        else:
            text = f"{text.capitalize()}. "
            print(chunk_filename, ":", text)
            whole_text += text
    # return the text for all chunks detected
    return whole_text

# Function to perform Normalization
def normalization(text):
    """
    Normalize the input text by:
    - Removing unwanted characters (except valid Markdown symbols)
    - Normalizing Unicode characters (e.g., é → e)
    - Optionally lowercase (commented out)
    """

    # Step 1: Normalize unicode (e.g., é → e +  ́)
    text = unicodedata.normalize('NFKD', text)

    # Step 2: Encode to ASCII and ignore characters like emojis
    text = text.encode('ASCII', 'ignore').decode()

    # Step 3: Remove unwanted characters, keeping markdown formatting
    text = re.sub(r"[^A-Za-z0-9\s\.\,\!\?\:\-\*\_\#\|\[\]\(\)]", '', text)

    # Step 4: Optional: lowercase (uncomment if needed)
    text = text.lower()

    return text